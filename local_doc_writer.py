import os
from datetime import datetime
from typing import Dict, Tuple

from docx import Document
from docx.shared import Pt

from config import DOCUMENTS_OUTPUT_DIR
from tools import slugify_filename


def ensure_output_dir():
    folder = os.path.abspath(DOCUMENTS_OUTPUT_DIR)
    os.makedirs(folder, exist_ok=True)
    return folder


def _looks_numbered(text: str) -> bool:
    if len(text) < 3:
        return False
    parts = text.split(". ", 1)
    if len(parts) != 2:
        return False
    return parts[0].isdigit()


def _remove_number_prefix(text: str) -> str:
    parts = text.split(". ", 1)
    if len(parts) == 2 and parts[0].isdigit():
        return parts[1].strip()
    return text.strip()


def parse_structured_document_text(title: str, body_text: str):
    elements = []

    for raw_line in body_text.splitlines():
        line = raw_line.strip()

        if not line:
            elements.append({"type": "blank"})
            continue

        if line.endswith(":") and not line.startswith("- ") and not _looks_numbered(line):
            elements.append({"type": "heading", "text": line[:-1].strip()})
            continue

        if line.startswith("- "):
            elements.append({"type": "bullet", "text": line[2:].strip()})
            continue

        if _looks_numbered(line):
            elements.append({"type": "number", "text": _remove_number_prefix(line)})
            continue

        elements.append({"type": "paragraph", "text": line})

    return {
        "title": title.strip() or "Eve Document",
        "elements": elements,
    }


def build_docx_file(title: str, body_text: str) -> Tuple[bool, str, Dict]:
    try:
        output_dir = ensure_output_dir()
        parsed = parse_structured_document_text(title, body_text)

        doc = Document()

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(parsed["title"])
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_para.style = doc.styles["Title"]

        for element in parsed["elements"]:
            kind = element["type"]

            if kind == "blank":
                doc.add_paragraph("")
            elif kind == "heading":
                doc.add_heading(element["text"], level=1)
            elif kind == "bullet":
                doc.add_paragraph(element["text"], style="List Bullet")
            elif kind == "number":
                doc.add_paragraph(element["text"], style="List Number")
            elif kind == "paragraph":
                doc.add_paragraph(element["text"])

        safe_name = slugify_filename(parsed["title"], fallback="eve_document")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(output_dir, f"{safe_name}_{timestamp}.docx")

        doc.save(file_path)

        if not os.path.exists(file_path):
            return False, "The file save completed but the document was not found on disk.", {}

        info = {
            "title": parsed["title"],
            "path": os.path.abspath(file_path),
            "filename": os.path.basename(file_path),
        }
        return True, "Document created successfully.", info

    except Exception as e:
        return False, f"Failed to create local document: {e}", {}